import json
import re
import traceback
import docx
from docx.oxml.table import CT_Tbl
from sentence_transformers import SentenceTransformer, util
import numpy as np
import sys
import faiss


# Load a SentenceTransformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

def is_table(paragraph):
    return any(isinstance(element, CT_Tbl) for element in paragraph._element)

def add_subheadings(paragraphs, index, level):
    subheadings = {}
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            para_level = int(paragraph.style.name.split()[-1])
            if para_level == level:
                index += 1
                subheading = {
                    'text': get_text_for_section(paragraphs, index),
                    'subheadings': add_subheadings_with_text(paragraphs, index, level + 1),
                    'bullet_points': get_bullet_points(paragraphs, index + 1)
                }
                subheadings[paragraph.text] = subheading
            else:
                break
        else:
            index += 1
    return subheadings

def add_subheadings_with_text(paragraphs, index, level):
    subheadings = {}
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            para_level = int(paragraph.style.name.split()[-1])
            if para_level == level:
                index += 1
                subheading = {
                    'text': get_text_for_section(paragraphs, index),
                    'subheadings': add_subheadings_with_text(paragraphs, index, level + 1),
                    'bullet_points': get_bullet_points(paragraphs, index + 1)
                }
                subheadings[paragraph.text] = subheading
            else:
                break
        else:
            index += 1
    return subheadings


def get_text_for_section(paragraphs, index):
    section_text = []
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            break
        elif is_table(paragraph):
            section_text.append(get_text_from_table(paragraph))
        else:
            section_text.append(paragraph.text)
        index += 1
    return '\n'.join(section_text)


def get_text_from_table(table_paragraph):
    table = table_paragraph._element
    text = ''
    for row in table.iter_tbl():
        for cell in row.iter_tc():
            for p in cell.iter_p():
                text += ' '.join(run.text for run in p.iter_t())
                text += '\n'
    return text

def get_headings_and_subheadings_with_text(doc):
    headings = {}
    index = 0
    current_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            if level == 1:
                current_heading = paragraph.text
                heading = {
                    'text': get_text_for_section(doc.paragraphs, index + 1),
                    'subheadings': add_subheadings(doc.paragraphs, index + 1, level + 1),
                    # Added bullet points
                    'bullet_points': get_bullet_points(doc.paragraphs, index + 1)
                }
                headings[current_heading] = heading
        index += 1
    return headings

def get_headings_and_subheadings_with_text_with_minlevel(doc, min_level=1):
    headings = {}
    index = 0
    current_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            if level < min_level:
                continue
            if level >= min_level:
                current_heading = paragraph.text
                heading = {
                    'text': get_text_for_section(doc.paragraphs, index + min_level),
                    'subheadings': add_subheadings(doc.paragraphs, index + 1, level + 1),
                    # Added bullet points
                    'bullet_points': get_bullet_points(doc.paragraphs, index + min_level)
                }
                headings[current_heading] = heading
        index += 1
    return headings
    
def preprocess_text(text):
    # Tokenization, lowercase, and remove punctuation
    # You might want to enhance this preprocessing step
    return text

def calculate_cosine_similarity(embeddings1, embeddings2):
    # Calculate cosine similarity between two sets of embeddings
    return util.pytorch_cos_sim(embeddings1, embeddings2)

def get_combined_text(heading, data):
    combined_text = f"Title: {heading} : {data['text']}\n"
    for subheading, subheading_data in data['subheadings'].items():
        subheading_text = subheading_data['text']
        combined_text += f"Subsection: {subheading} - {subheading_text}\n"

    return combined_text

################### Bullet Points ############################

def get_bullet_points(paragraphs, index):
    bullet_points = []
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            break
        elif any(paragraph.style.name.startswith(prefix) for prefix in ['ListBullet', 'ListNumber', 'ListParagraph']):
            bullet_points.append(f"- {paragraph.text}")  # Include bullet point text
        index += 1
    # return '\n'.join(bullet_points)
    return bullet_points

def get_bullets_embeddings(doc, filter_subheadings=False):
    doc_bullets = []
    for heading, data in doc.items():
        if filter_subheadings:
            if re.search("INTRODUCTION|OBJECTIVE|METHODS", heading, flags=re.IGNORECASE):
                for subheading, sub_data in data['subheadings'].items():
                   for bullet in sub_data['bullet_points']:
                    doc_bullets.append(**bullet)
        else:
            for subheading, sub_data in data['subheadings'].items():
                for bullet in sub_data['bullet_points']:
                    doc_bullets.append(**bullet)
        
    doc_bullets_preprocessed, doc_bullets_embeddings = preprocess_and_encode(doc_bullets, model)

    return doc_bullets, doc_bullets_preprocessed, doc_bullets_embeddings

def get_best_bullet(doc1_bullets, doc1_bullets_emb, doc2_bullets, bullets_index):
    bullets_match = []

    for i in range(len(doc1_bullets)):
        doc1_bullet_text = doc1_bullets[i]
        doc1_bullets_embedding = doc1_bullets_emb[i]
        
        # Retrieve similar sub heading from doc2
        # similar_indices = retrieve_best_match(doc1_subheading_embedding, subheads_index)
        score, similar_indices = retrieve_best_match(doc1_bullets_embedding, bullets_index)
        
        best_match_bullets = doc2_bullets[similar_indices]
        best_match_data = {
            'doc1_bullets': doc1_bullet_text,
            'best_match_text': best_match_bullets,
            'similar_indices' : similar_indices,
            'score' : score,
            'match_type' : get_match_type(score),
        }
        print("bullets_match : ", best_match_data)
        bullets_match.append(best_match_data)
            
    return bullets_match



########### EMBEDDING MATCHING ######################

def get_match_type(score):
    score = float(score)

    if score >= 0.8:
        match_type = "Full Match"
    elif score >= 0.5:
        match_type = "Partial Match"
    else:
        match_type = "No Match"
    
    return match_type


def preprocess_and_encode(texts, model):
    preprocessed_texts = [preprocess_text(text) for text in texts]
    embeddings = model.encode(preprocessed_texts, convert_to_tensor=True)
    return preprocessed_texts, embeddings


def build_index(embeddings):
    print(embeddings.shape)
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)  # Inner product similarity
    # faiss.normalize_L2(embeddings)  # Normalize embeddings
    
    # index = faiss.IndexFlatL2(dimension)  # L2 distance for cosine similarity
    index.add(embeddings)
    return index

# def retrieve_best_match(query_embedding, index, k=1):
#     _, indices = index.search(query_embedding, k)
#     return indices[0]
def retrieve_best_match(query_embedding, index, k=3):
    # scores, indices = index.search(query_embedding, k)
    query_embedding_2d = np.expand_dims(query_embedding, axis=0)  # Reshape to 2D
    scores, indices = index.search(query_embedding_2d, k)
    # print("scores : ", scores)
    # print("indices : ", indices)
    # return scores[0][0], indices[0][0]

    # Transform scores to fall within the range 0 to 1 using sigmoid function
    transformed_scores = 1 / (1 + np.exp(-scores))
    
    # print("scores : ", scores)
    # print("transformed_scores : ", transformed_scores)
    # print("indices : ", indices)
    return transformed_scores[0][0], indices[0][0]


def get_headings_embeddings(doc):
    doc_headings = [get_combined_text(heading, data) for heading, data in doc.items()]
    doc_headings_preprocessed, doc_headings_embeddings = preprocess_and_encode(doc_headings, model)

    return doc_headings, doc_headings_preprocessed, doc_headings_embeddings


def get_subheadings_embeddings(doc, filter_subheadings=False):
    doc_subheadings = []
    for heading, data in doc.items():
        if filter_subheadings:
            if re.search("INTRODUCTION|OBJECTIVE|METHODS", heading, flags=re.IGNORECASE):
                for subheading, sub_data in data['subheadings'].items():
                    doc_subheadings.append(get_combined_text(subheading, sub_data))
        else:
            for subheading, sub_data in data['subheadings'].items():
                    doc_subheadings.append(get_combined_text(subheading, sub_data))
        
    doc_subheadings_preprocessed, doc_subheadings_embeddings = preprocess_and_encode(doc_subheadings, model)

    return doc_subheadings, doc_subheadings_preprocessed, doc_subheadings_embeddings


def get_best_method(doc1_subheads, doc1_subheads_emb, doc2_subheads, subheads_index):
    methods_match = []
    for i in range(len(doc1_subheads)):
        doc1_subheading_text = doc1_subheads[i]
        doc1_subheading_embedding = doc1_subheads_emb[i]

        # Retrieve similar sub heading from doc2
        # similar_indices = retrieve_best_match(doc1_subheading_embedding, subheads_index)
        score, similar_indices = retrieve_best_match(doc1_subheading_embedding, subheads_index)
        best_match_method = doc2_subheads[similar_indices]
        best_match_data = {
            'doc1_subheading_text': doc1_subheading_text,
            'best_match_text': best_match_method,
            'similar_indices' : similar_indices,
            'score' : score,
            'match_type' : get_match_type(score),
        }
        print("best_match_data METHODS : ", best_match_data)
        methods_match.append(best_match_data)
    
    return methods_match

def get_best_intro_obj(doc1_heads, doc1_heads_emb, doc2_heads, heads_index):
    headings_match = []

    for i in range(len(doc1_heads)):
        doc1_heading_text = doc1_heads[i]
        doc1_heading_embedding = doc1_heads_emb[i]
        if re.search("Title: INTRODUCTION|Title: OBJECTIVE", doc1_heading_text, flags=re.IGNORECASE):
            # Retrieve similar sub heading from doc2
            # similar_indices = retrieve_best_match(doc1_subheading_embedding, subheads_index)
            score, similar_indices = retrieve_best_match(doc1_heading_embedding, heads_index)
            
            best_match_heading = doc2_heads[similar_indices]
            best_match_data = {
                'doc1_heading_text': doc1_heading_text,
                'best_match_text': best_match_heading,
                'similar_indices' : similar_indices,
                'score' : score,
                'match_type' : get_match_type(score),
            }
            print("best_match_data INTRODUCTION OR OBJECTIVES : ", best_match_data)
            headings_match.append(best_match_data)
        else:
            pass
    
    return headings_match


def main():
    
    doc1_path = '/CSR_Sample_1.docx'
    doc2_path = '/Protocol_Sample_2.docx'
    
    try:
        doc1 = docx.Document(doc1_path)
        doc2 = docx.Document(doc2_path)

        headings_structure_doc1 = get_headings_and_subheadings_with_text(doc1)
        headings_structure_doc2 = get_headings_and_subheadings_with_text_with_minlevel(doc2, min_level=1)
        # headings_structure_doc2 = get_headings_and_subheadings_with_text1(doc2)
        

        # Get doc1 headings and sub headings embeddings
        doc1_heads, doc1_heads_preprocess, doc1_heads_emb = get_headings_embeddings(headings_structure_doc1)
        doc1_subheads, doc1_subheads_preprocess, doc1_subheads_emb = get_subheadings_embeddings(headings_structure_doc1, 
                                                                                                filter_subheadings=True)
        # Convert embeddings from tensors to NumPy arrays
        doc1_heads_emb = doc1_heads_emb.detach().numpy()
        doc1_subheads_emb = doc1_subheads_emb.detach().numpy()



        # Get doc2 headings and sub headings embeddings
        doc2_heads, doc2_heads_preprocess, doc2_heads_emb = get_headings_embeddings(headings_structure_doc2)
        doc2_subheads, doc2_subheads_preprocess, doc2_subheads_emb = get_subheadings_embeddings(headings_structure_doc2,
                                                                                                filter_subheadings=False)
        # Convert embeddings from tensors to NumPy arrays
        doc2_heads_emb = doc2_heads_emb.detach().numpy()
        doc2_subheads_emb = doc2_subheads_emb.detach().numpy()


        # Bullet Points
        doc1_bullets, doc1_bullets_preprocessed, doc1_bullets_emb = get_bullets_embeddings(headings_structure_doc1, 
                                                                                               filter_subheadings=False)
        doc2_bullets, doc2_bullets_preprocessed, doc2_bullets_emb = get_bullets_embeddings(headings_structure_doc2, 
                                                                                               filter_subheadings=False)
        # Convert embeddings from tensors to NumPy arrays
        doc1_bullets_emb = doc1_bullets_emb.detach().numpy()
        doc2_bullets_emb = doc2_bullets_emb.detach().numpy()


        
        # Build the index using the doc2 embeddings for main headings and sub headings and bullets
        heads_index = build_index(doc2_heads_emb)
        subheads_index = build_index(doc2_subheads_emb)
        bullets_index = build_index(doc2_bullets_emb)

        # print("##############################")

        # for heading, data in headings_structure_doc1.items():
        #     print(f"Heading: {heading}")
        #     subheadings = data['subheadings']
        #     for subheading in subheadings.keys():
        #         print(f"  Subheading: {subheading}")

        # print("##############################")
        # Matching Logic
        all_headings_match = []
        all_methods_match = []
        all_bullets_match = []
        
        for doc1_heading, doc1_data in headings_structure_doc1.items():
            if re.search("METHODS", doc1_heading, flags=re.IGNORECASE):
                methods_match = get_best_method(doc1_subheads, doc1_subheads_emb, doc2_subheads, subheads_index)
                all_methods_match = all_methods_match + methods_match

                bullets_match = get_best_bullet(doc1_bullets, doc1_bullets_emb, doc2_bullets, bullets_index)
                all_bullets_match = all_bullets_match + bullets_match
                
            
            elif re.search("INTRODUCTION|OBJECTIVE", doc1_heading, flags=re.IGNORECASE):
                headings_match = get_best_intro_obj(doc1_heads, doc1_heads_emb, doc2_heads, heads_index)
                all_headings_match = all_headings_match + headings_match
            # if re.search("INTRODUCTION|OBJECTIVE", doc1_heading, flags=re.IGNORECASE):
            #     headings_match = get_best_intro_obj(doc1_heads, doc1_heads_emb, doc2_heads, heads_index)
            #     all_headings_match = all_headings_match + headings_match
            else:
                pass
        
        # Add matched results to the "matches" dictionary
        print("all_headings_match : ", all_headings_match)
        print("###################################################")
        print("all_methods_match : ", all_methods_match)
        print("###################################################")
        print("bullets_match : ", bullets_match)

    except Exception as e:
        err = traceback.format_exc()
        print(f"An error occurred: {e}, {err}")

if __name__ == '__main__':
    main()
