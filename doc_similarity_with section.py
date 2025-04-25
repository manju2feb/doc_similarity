import json
import re
import traceback
import docx
from docx.oxml.table import CT_Tbl
from sentence_transformers import SentenceTransformer, util
import numpy as np
import sys


# Load a SentenceTransformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

def is_table(paragraph):
    return any(isinstance(element, CT_Tbl) for element in paragraph._element)

def add_subheadings(paragraphs, index, level):
    subheadings = {}
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            para_level = int(paragraph.style.name.split()[-1])
            if para_level == level:
                index += 1
                subheading = {
                    'text': get_text_for_section(paragraphs, index),
                    'subheadings': add_subheadings(paragraphs, index, level + 1)
                }
                subheadings[paragraph.text] = subheading
            else:
                break
        else:
            index += 1
    return subheadings

def add_subheadings_with_text(paragraphs, index, level):
    subheadings = {}
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.style.name.startswith('Heading'):
            para_level = int(paragraph.style.name.split()[-1])
            if para_level == level:
                index += 1
                subheading = {
                    'text': get_text_for_section(paragraphs, index),
                    'subheadings': add_subheadings_with_text(paragraphs, index, level + 1)
                }
                subheadings[paragraph.text] = subheading
            else:
                break
        else:
            index += 1
    return subheadings

def get_text_for_section(paragraphs, index):
    section_text = []
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        print("paragraph : ", paragraph)
        if paragraph.style.name.startswith('Heading'):
            break
        elif is_table(paragraph):
            section_text.append(get_text_from_table(paragraph))
        # elif paragraph.style.name.contains('Bullet'):
        elif paragraph.style.name.startswith('ListBullet') or paragraph.style.name.startswith('ListBullet') or paragraph.style.name.startswith('ListNumber') or 'Number' in paragraph.style.name:
            print("paragraph : bullet", paragraph)
            section_text.append(f"- {paragraph.text}")  # Include bullet point text
        else:
            section_text.append(paragraph.text)
        index += 1
    return '\n'.join(section_text)

def get_text_from_table(table_paragraph):
    table = table_paragraph._element
    text = ''
    for row in table.iter_tbl():
        for cell in row.iter_tc():
            for p in cell.iter_p():
                text += ' '.join(run.text for run in p.iter_t())
                text += '\n'
    return text

def get_headings_and_subheadings_with_text(doc):
    headings = {}
    index = 0
    current_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            if level == 1:
                current_heading = paragraph.text
                heading = {
                    'text': get_text_for_section(doc.paragraphs, index),
                    'subheadings': add_subheadings(doc.paragraphs, index + 1, level + 1)
                }
                headings[current_heading] = heading
        index += 1
    return headings

def get_headings_and_subheadings_with_text_with_minlevel(doc, min_level=1):
    headings = {}
    index = 0
    current_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            if level < min_level:
                continue
            if level >= min_level:
                current_heading = paragraph.text
                heading = {
                    'text': get_text_for_section(doc.paragraphs, index),
                    'subheadings': add_subheadings(doc.paragraphs, index + 1, level + 1)
                }
                headings[current_heading] = heading
        index += 1
    return headings

    
def preprocess_text(text):
    # Tokenization, lowercase, and remove punctuation
    # You might want to enhance this preprocessing step
    return text

def calculate_cosine_similarity(embeddings1, embeddings2):
    # Calculate cosine similarity between two sets of embeddings
    return util.pytorch_cos_sim(embeddings1, embeddings2)

# def find_most_similar_text(target_text, candidate_texts):
#     target_embedding = model.encode([preprocess_text(target_text)], convert_to_tensor=True)
#     candidate_embeddings = model.encode(candidate_texts, convert_to_tensor=True)
#     similarity_scores = np.array(util.pytorch_cos_sim(target_embedding, candidate_embeddings)[0])
#     best_match_index = np.argmax(similarity_scores)
#     return candidate_texts[best_match_index]

def get_combined_text(heading, data):
    combined_text = f"Title: {heading} : {data['text']}\n"
    for subheading, subheading_data in data['subheadings'].items():
        subheading_text = subheading_data['text']
        combined_text += f"Subsection: {subheading} - {subheading_text}\n"

    return combined_text

def find_most_similar_text(target_text, candidate_texts):
    target_embedding = model.encode([preprocess_text(target_text)], convert_to_tensor=True)
    candidate_embeddings = model.encode([preprocess_text(text) for text in candidate_texts], convert_to_tensor=True)
    similarity_scores = np.array(util.pytorch_cos_sim(target_embedding, candidate_embeddings)[0])
    best_match_index = np.argmax(similarity_scores)
    return candidate_texts[best_match_index]

def find_most_similar_text(target_text, candidate_texts):
    target_embedding = model.encode([preprocess_text(target_text)], convert_to_tensor=True)
    # Preprocess and filter out empty texts from candidate_texts
    preprocessed_candidate_texts = [preprocess_text(text) for text in candidate_texts if preprocess_text(text)]
    if not preprocessed_candidate_texts:
        return None # Return early if all candidate texts became empty after preprocessing

    candidate_embeddings = model.encode(preprocessed_candidate_texts, convert_to_tensor=True)
    similarity_scores = np.array(util.pytorch_cos_sim(target_embedding, candidate_embeddings)[0])
    best_match_index = np.argmax(similarity_scores)

    if best_match_index < len(candidate_texts):
        return candidate_texts[best_match_index]
    else:
        return None
    
def match_with_sap(doc1_heading, doc2_heading, doc1_heading_text, doc2_heading_text, doc1_data, doc2_data):
    # Calculate cosine similarity for headings
    heading_similarity_scores = calculate_cosine_similarity(
        model.encode([doc1_heading_text], convert_to_tensor=True),
        model.encode([doc2_heading_text], convert_to_tensor=True)
    )

    
    
    return heading_similarity_scores

def best_match_found(best_match):
    
    headings_best_match_data = {}
    # If best_match is found for a given heading
    if best_match:
        best_match_headings_csr_heading = best_match['csr_heading']
        best_match_headings_sap_heading = best_match['sap_heading']
        best_match_headings_csr_heading_text = best_match['csr_heading_combined_text']
        best_match_headings_sap_heading_text = best_match['sap_heading_combined_text']
        best_match_headings_score = best_match['best_match_score']
        csr_heading_subheadings = best_match['csr_heading_subheadings']
        sap_heading_subheadings = best_match['sap_heading_subheadings']

        headings_best_match_data =  {
                'csr_heading': best_match_headings_csr_heading,
                'sap_heading': best_match_headings_sap_heading,
                'csr_heading_combined_text': best_match_headings_csr_heading_text,
                'sap_heading_combined_text': best_match_headings_sap_heading_text,
                'score': best_match_headings_score,
                'csr_heading_subheadings' : csr_heading_subheadings,
                'sap_heading_subheadings' : sap_heading_subheadings,
            }
        
        matched_subsection = find_most_similar_text(
            best_match_headings_csr_heading_text,
            [subheading_data['text'] for subheading_data in sap_heading_subheadings.values()]
        )
        headings_best_match_data['matched_subsection'] = matched_subsection
        if headings_best_match_data['score'] >= 0.8:
            headings_best_match_data['match_type'] = "Full Match"
        elif headings_best_match_data['score'] >= 0.5:
            headings_best_match_data['match_type'] = "Partial Match"
        else:
            headings_best_match_data['match_type'] = "No Match"
    
    return headings_best_match_data
                


def main():
   
    doc1_path = '/UW/AK/CSR_Sample.docx'
    doc2_path = '/UW/AK/Protocol_Sample_1.docx'
    
    try:
        doc1 = docx.Document(doc1_path)
        doc2 = docx.Document(doc2_path)

        headings_structure_doc1 = get_headings_and_subheadings_with_text(doc1)
        headings_structure_doc2 = get_headings_and_subheadings_with_text_with_minlevel(doc2, min_level=1)

        # Pretty print the dictionary
        pretty_data = json.dumps(headings_structure_doc1, indent=4)
        print(headings_structure_doc1)

        headings_match = []
        for doc1_heading, doc1_data in headings_structure_doc1.items():
            if not re.search("INTRODUCTION|OBJECTIVE|METHODS", doc1_heading, flags=re.IGNORECASE):
                continue

            best_match_score = 0.0
            best_match = None

            doc1_heading_text = preprocess_text(get_combined_text(doc1_heading, doc1_data))
            
            if re.search("METHODS", doc1_heading, flags=re.IGNORECASE):
                for doc2_heading, doc2_data in headings_structure_doc2.items():
                    for doc2_subheading, doc2_subdata in doc2_data['subheadings'].items():
                        # doc2_heading_text = preprocess_text(doc2_data['text'])
                        doc2_heading_text = preprocess_text(get_combined_text(doc2_subheading, doc2_subdata))
                        
                        print("doc2_heading_text : ", doc2_heading_text)
                        print("###############################")
                    
                        heading_similarity_scores = match_with_sap(doc1_heading, doc2_heading, 
                                                        doc1_heading_text, doc2_heading_text, doc1_data, doc2_data)
                        
                        # Update best match if current score is higher
                        if heading_similarity_scores > best_match_score:
                            best_match_score = heading_similarity_scores
                            # best_match = (doc1_heading, doc2_heading, doc1_heading_text, doc2_heading_text)
                            # Create a best_match dictionary containing match information
                            best_match = {
                                'csr_heading': doc1_heading,
                                'sap_heading': doc2_heading,
                                'csr_heading_combined_text': doc1_heading_text,
                                'sap_heading_combined_text': doc2_heading_text,
                                'best_match_score': best_match_score,
                                'csr_heading_subheadings': doc1_data['subheadings'],
                                'sap_heading_subheadings': doc2_data['subheadings']
                            }
                    
                if best_match:
                    headings_best_match_data = best_match_found(best_match)
                    if headings_best_match_data != {}:
                        headings_match.append(headings_best_match_data)
            else:
                for doc2_heading, doc2_data in headings_structure_doc2.items():
                    # doc2_heading_text = preprocess_text(doc2_data['text'])
                    doc2_heading_text = preprocess_text(get_combined_text(doc2_heading, doc2_data))
                    
                    print("doc2_heading_text : ", doc2_heading_text)
                    print("###############################")
                
                    heading_similarity_scores = match_with_sap(doc1_heading, doc2_heading, 
                                                      doc1_heading_text, doc2_heading_text, doc1_data, doc2_data)
                    
                    # Update best match if current score is higher
                    if heading_similarity_scores > best_match_score:
                        best_match_score = heading_similarity_scores
                        # best_match = (doc1_heading, doc2_heading, doc1_heading_text, doc2_heading_text)
                        # Create a best_match dictionary containing match information
                        best_match = {
                            'csr_heading': doc1_heading,
                            'sap_heading': doc2_heading,
                            'csr_heading_combined_text': doc1_heading_text,
                            'sap_heading_combined_text': doc2_heading_text,
                            'best_match_score': best_match_score,
                            'csr_heading_subheadings': doc1_data['subheadings'],
                            'sap_heading_subheadings': doc2_data['subheadings']
                        }
                
                if best_match:
                    headings_best_match_data = best_match_found(best_match)
                    if headings_best_match_data != {}:
                        headings_match.append(headings_best_match_data)

            
            
        
        # Add matched results to the "matches" dictionary
        print("headings_match : ", headings_match)

    except Exception as e:
        err = traceback.format_exc()
        print(f"An error occurred: {e}, {err}")

if __name__ == '__main__':
    main()